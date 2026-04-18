import os
import wikipedia
from docx import Document
from fpdf import FPDF

data_dir = "data"
os.makedirs(data_dir, exist_ok=True)

topics = [
    "History of Pakistan",
    "Geography of Pakistan",
    "Culture of Pakistan",
    "Economy of Pakistan",
    "Education in Pakistan",
    "Tourism in Pakistan",
    "Languages of Pakistan",
    "Demographics of Pakistan",
    "Politics of Pakistan"
]

# -------------------
# TXT FILES
# -------------------

for i, topic in enumerate(topics[:10], 1):
    text = wikipedia.page(topic).content[:3000]  # first 3000 characters
    
    with open(os.path.join(data_dir, f"topic_{i}.txt"), "w", encoding="utf-8") as f:
        f.write(text)

# # -------------------
# # DOCX FILES
# # -------------------

for i, topic in enumerate(topics[:5], 1):
    text = wikipedia.page(topic).content[:3000]

    doc = Document()
    doc.add_heading(topic, 0)
    doc.add_paragraph(text)

    doc.save(os.path.join(data_dir, f"topic_{i}.docx"))

# -------------------
# PDF FILES
# -------------------

for i, topic in enumerate(topics[2:3], 1):
    text = wikipedia.page(topic).content[:3000]

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    pdf.cell(200, 10, topic, ln=True, align="C")
    pdf.ln(10)
    text = text.encode("latin-1", "replace").decode("latin-1")
    pdf.multi_cell(0, 10, text)

    pdf.output(os.path.join(data_dir, f"topic_{1}.pdf"))

print("Files created using internet content!")